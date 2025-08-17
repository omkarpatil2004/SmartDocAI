import streamlit as st
from utils.pdf_handler import extract_text_from_pdf
from utils.ocr_handler import extract_text_from_image
from utils.audio_handler import transcribe_audio
from utils.video_handler import transcribe_video
from utils.summarizer import abstractive_summary
from docx import Document
import tempfile
import os

# Helper function to save output
def save_to_docx(text, filename="output.docx"):
    doc = Document()
    doc.add_heading("Smart DocAI Notes", level=1)
    doc.add_paragraph(text)
    path = os.path.join("outputs", filename)
    doc.save(path)
    return path

# Streamlit UI
st.set_page_config(page_title="Smart DocAI", layout="centered")
st.title("📑 Smart DocAI")
st.write("Convert Video, Audio, PDF, and Images into Notes (.docx)")

option = st.selectbox("Choose an option:", 
                      ["📹 Video to Notes", "🎤 Audio to Notes", "📄 Image/PDF/Doc to Notes"])

mode = st.radio("Choose mode:", ["As-it-is", "Summarization"])

uploaded_file = st.file_uploader("Upload your file", type=["mp4", "mp3", "wav", "pdf", "png", "jpg", "jpeg", "docx"])

if uploaded_file is not None:
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        tmp_file.write(uploaded_file.read())
        file_path = tmp_file.name

    extracted_text = ""

    if option == "📹 Video to Notes":
        extracted_text = transcribe_video(file_path)

    elif option == "🎤 Audio to Notes":
        extracted_text = transcribe_audio(file_path)

    elif option == "📄 Image/PDF/Doc to Notes":
        if uploaded_file.name.endswith(".pdf"):
            extracted_text = extract_text_from_pdf(file_path)
        else:
            extracted_text = extract_text_from_image(file_path)

    if mode == "Summarization":
        extracted_text = abstractive_summary(extracted_text)

    if extracted_text:
        output_path = save_to_docx(extracted_text, "output.docx")
        with open(output_path, "rb") as f:
            st.download_button("Download Notes (.docx)", f, file_name="SmartDocAI_Notes.docx")
